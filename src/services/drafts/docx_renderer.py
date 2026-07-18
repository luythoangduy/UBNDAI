"""Sinh DOCX deterministic cho bản nháp kết quả thủ tục.

Mỗi biểu mẫu có thông số vật lý trong manifest. Giấy khai sinh được dựng thành
hai trang tương ứng mặt trước và mặt sau của phôi; file chỉ phục vụ rà soát và
không thay thế phôi do Bộ Tư pháp phát hành.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from io import BytesIO
from typing import Callable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from src.models import DraftGenerateRequest, DraftTemplate, GeneratedDraft
from src.services.drafts.registry import get_template
from src.services.drafts.renderer import WATERMARK, _render_block, generate

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@dataclass(frozen=True)
class GeneratedDocx:
    content: bytes
    filename: str
    media_type: str = DOCX_MEDIA_TYPE


def generate_docx(payload: DraftGenerateRequest) -> GeneratedDocx:
    """Validate dữ liệu bằng renderer chung rồi dựng file Word tương ứng."""

    result = generate(payload)
    template = get_template(payload.procedure_id, payload.template_id)
    builder = _BUILDERS.get(template.id)
    if builder is None:
        builder = _build_generic_template
    document = builder(template, result)
    output = BytesIO()
    document.save(output)
    return GeneratedDocx(
        content=output.getvalue(),
        filename=template.docx_style.filename,
    )


def _build_generic_template(
    template: DraftTemplate, result: GeneratedDraft
) -> DocxDocument:
    """Render any declarative template without a procedure-specific builder.

    Specialized builders may still encode regulated physical forms, while this
    fallback produces a clearly watermarked review draft from the manifest.
    """
    doc = Document()
    _configure_document(doc, template)
    for block in template.layout:
        text = _render_block(block, template, result.normalized_values)
        if block.kind == "title":
            _add_title(doc, text, template)
            continue
        paragraph = doc.add_paragraph()
        _format_body_paragraph(paragraph, template)
        if block.kind != "spacer":
            run = paragraph.add_run(text)
            _format_run(
                run,
                template.docx_style.body_font,
                template.docx_style.body_size_pt,
            )
    disclaimer = doc.add_paragraph()
    _format_body_paragraph(disclaimer, template)
    run = disclaimer.add_run(template.disclaimer)
    _format_run(
        run,
        template.docx_style.body_font,
        template.docx_style.notes_font_size_pt,
        italic=True,
        color="777777",
    )
    return doc


def _build_birth_certificate(
    template: DraftTemplate, result: GeneratedDraft
) -> DocxDocument:
    doc = Document()
    _configure_document(doc, template)
    values = result.normalized_values

    _add_centered_text(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True)
    _add_centered_text(doc, "Độc lập - Tự do - Hạnh phúc", bold=True)
    _add_emblem_placeholder(doc, template)
    _add_title(doc, "GIẤY KHAI SINH", template)

    reference = _new_fixed_table(doc, [82, 82], hidden_borders=True)
    _set_cell_line(
        reference.cell(0, 0),
        "Số",
        _value(values, "so", "Số"),
        template,
    )
    _set_cell_line(
        reference.cell(0, 1),
        "Quyển số",
        _value(values, "quyen_so", "Quyển số"),
        template,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    _add_field_line(
        doc,
        "Họ, chữ đệm, tên",
        _value(values, "ho_ten_con", "Họ, chữ đệm, tên"),
        template,
        value_bold=True,
    )
    _add_field_line(
        doc,
        "Ngày, tháng, năm sinh (bằng số)",
        _date_value(values, "ngay_sinh", "Ngày, tháng, năm sinh", words=False),
        template,
    )
    _add_field_line(
        doc,
        "Ngày, tháng, năm sinh (bằng chữ)",
        _date_value(values, "ngay_sinh", "Ngày, tháng, năm sinh", words=True),
        template,
    )
    _add_value_row(
        doc,
        [("Giới tính", "gioi_tinh"), ("Dân tộc", "dan_toc"), ("Quốc tịch", "quoc_tich")],
        [47, 52, 65],
        values,
        template,
    )
    _add_field_line(doc, "Nơi sinh", _value(values, "noi_sinh", "Nơi sinh"), template)
    _add_field_line(doc, "Quê quán", _value(values, "que_quan", "Quê quán"), template)
    _add_field_line(
        doc,
        "Số định danh cá nhân",
        _value(values, "so_dinh_danh_ca_nhan", "Số định danh cá nhân"),
        template,
    )

    _add_parent_block(doc, "me", values, template)
    _add_parent_block(doc, "cha", values, template)

    _add_field_line(
        doc,
        "Họ, chữ đệm, tên người đi đăng ký khai sinh",
        _value(
            values,
            "ho_ten_nguoi_di_dang_ky",
            "Họ, chữ đệm, tên người đi đăng ký khai sinh",
        ),
        template,
    )
    _add_field_line(
        doc,
        "Giấy tờ tùy thân",
        _value(
            values,
            "giay_to_tuy_than_nguoi_di_dang_ky",
            "Giấy tờ tùy thân của người đi đăng ký khai sinh",
        ),
        template,
    )
    _add_field_line(
        doc,
        "Nơi đăng ký khai sinh",
        _value(values, "noi_dang_ky", "Nơi đăng ký khai sinh"),
        template,
    )
    _add_field_line(
        doc,
        "Ngày đăng ký",
        _date_value(values, "ngay_dang_ky", "Ngày đăng ký", words=False),
        template,
    )
    _add_signature_block(doc, values, template)
    _add_reverse_notes_page(doc, template)
    return doc


def _configure_document(doc: DocxDocument, template: DraftTemplate) -> None:
    style = template.docx_style
    section = doc.sections[0]
    section.page_width = Mm(style.page_width_mm)
    section.page_height = Mm(style.page_height_mm)
    section.top_margin = Mm(style.margin_top_mm)
    section.right_margin = Mm(style.margin_right_mm)
    section.bottom_margin = Mm(style.margin_bottom_mm)
    section.left_margin = Mm(style.margin_left_mm)
    section.header_distance = Mm(3)
    section.footer_distance = Mm(6)

    normal = doc.styles["Normal"]
    normal.font.name = style.body_font
    normal.font.size = Pt(style.body_size_pt)
    _set_style_fonts(normal._element, style.body_font)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(style.line_spacing_pt)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.line_spacing = Pt(9)
    run = header.add_run(WATERMARK)
    _format_run(run, style.body_font, 8, bold=True, color="A00000")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    footer.paragraph_format.line_spacing = Pt(8)
    run = footer.add_run("Bản nháp rà soát - không thay thế phôi hộ tịch hợp lệ")
    _format_run(run, style.body_font, 7, italic=True, color="666666")

    core = doc.core_properties
    core.title = f"{template.output_name} - Dự thảo"
    core.subject = f"Mẫu {template.version}"
    core.author = "TTHC Assist"
    core.comments = template.disclaimer


def _add_centered_text(doc: DocxDocument, text: str, *, bold: bool) -> None:
    template_font = doc.styles["Normal"].font.name or "Times New Roman"
    template_size = doc.styles["Normal"].font.size.pt
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    _set_paragraph_spacing(paragraph, line_pt=template_size + 3)
    run = paragraph.add_run(text)
    _format_run(run, template_font, template_size, bold=bold)


def _add_emblem_placeholder(doc: DocxDocument, template: DraftTemplate) -> None:
    table = _new_fixed_table(doc, [20], hidden_borders=False, total_width_mm=20)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Mm(20)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(paragraph, line_pt=7)
    run = paragraph.add_run("VỊ TRÍ\nQUỐC HUY\n20 × 20 mm")
    _format_run(run, template.docx_style.body_font, 5.5, bold=True, color="777777")
    after = doc.add_paragraph()
    _set_paragraph_spacing(after, line_pt=3)


def _add_title(doc: DocxDocument, text: str, template: DraftTemplate) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = Pt(template.docx_style.title_size_pt + 2)
    run = paragraph.add_run(text)
    _format_run(
        run,
        template.docx_style.body_font,
        template.docx_style.title_size_pt,
        bold=True,
        color=template.docx_style.title_color_hex,
    )


def _add_field_line(
    doc: DocxDocument,
    label: str,
    value: str,
    template: DraftTemplate,
    *,
    value_bold: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    _format_body_paragraph(paragraph, template)
    _append_label_value(paragraph, label, value, template, value_bold=value_bold)


def _add_value_row(
    doc: DocxDocument,
    fields: list[tuple[str, str]],
    widths_mm: list[float],
    values: dict[str, str],
    template: DraftTemplate,
) -> None:
    table = _new_fixed_table(doc, widths_mm, hidden_borders=True)
    for cell, (label, key) in zip(table.rows[0].cells, fields, strict=True):
        _set_cell_line(cell, label, _value(values, key, label), template)


def _add_parent_block(
    doc: DocxDocument,
    role: str,
    values: dict[str, str],
    template: DraftTemplate,
) -> None:
    label_role = "mẹ" if role == "me" else "cha"
    _add_field_line(
        doc,
        f"Họ, chữ đệm, tên {label_role}",
        _value(values, f"ho_ten_{role}", f"Họ, chữ đệm, tên {label_role}"),
        template,
    )
    _add_value_row(
        doc,
        [
            ("Năm sinh", f"nam_sinh_{role}"),
            ("Dân tộc", f"dan_toc_{role}"),
            ("Quốc tịch", f"quoc_tich_{role}"),
        ],
        [47, 52, 65],
        values,
        template,
    )
    _add_field_line(
        doc,
        f"Nơi cư trú của {label_role}",
        _value(values, f"noi_cu_tru_{role}", f"Nơi cư trú của {label_role}"),
        template,
    )


def _add_signature_block(
    doc: DocxDocument, values: dict[str, str], template: DraftTemplate
) -> None:
    table = _new_fixed_table(doc, [82, 82], hidden_borders=True)
    left, right = table.rows[0].cells
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_body_paragraph(left_p, template)
    run = left_p.add_run("DẤU CỦA CƠ QUAN")
    _format_run(run, template.docx_style.body_font, 11, bold=True, color="777777")
    note = left.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(note, line_pt=12)
    run = note.add_run("[Không tạo trong bản nháp]")
    _format_run(run, template.docx_style.body_font, 9, italic=True, color="777777")

    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(right_p, line_pt=14)
    run = right_p.add_run("NGƯỜI KÝ GIẤY KHAI SINH")
    _format_run(run, template.docx_style.body_font, 11, bold=True)
    hint = right.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(hint, line_pt=12)
    run = hint.add_run("(Ký, ghi rõ họ tên, chức vụ)")
    _format_run(run, template.docx_style.body_font, 9, italic=True)
    spacer = right.add_paragraph()
    _set_paragraph_spacing(spacer, line_pt=26)
    name = right.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(name, line_pt=14)
    signer = values.get("ho_ten_nguoi_ky") or "[CHƯA CÓ: Họ tên người ký]"
    run = name.add_run(signer)
    _format_run(
        run,
        template.docx_style.body_font,
        11,
        bold=True,
        color="777777" if signer.startswith("[") else None,
    )
    position = values.get("chuc_vu_nguoi_ky")
    if position:
        position_p = right.add_paragraph()
        position_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(position_p, line_pt=13)
        run = position_p.add_run(position)
        _format_run(run, template.docx_style.body_font, 10)


def _add_reverse_notes_page(doc: DocxDocument, template: DraftTemplate) -> None:
    style = template.docx_style
    reverse = doc.add_section(WD_SECTION.NEW_PAGE)
    reverse.page_width = Mm(style.page_width_mm)
    reverse.page_height = Mm(style.page_height_mm)
    # Phụ lục quy định chính xác bảng mặt sau 158 x 260 mm. Lề riêng này
    # đặt bảng cân giữa A4 và chừa chỗ cho section mark bắt buộc của Word.
    reverse.top_margin = Mm(10)
    reverse.right_margin = Mm(26)
    reverse.bottom_margin = Mm(10)
    reverse.left_margin = Mm(26)
    reverse.header_distance = Mm(3)
    reverse.footer_distance = Mm(6)
    table = _new_fixed_table(
        doc,
        [style.notes_table_width_mm],
        hidden_borders=False,
        total_width_mm=style.notes_table_width_mm,
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Mm(style.notes_table_height_mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = Pt(16)
    run = paragraph.add_run("PHẦN GHI CHÚ NHỮNG THÔNG TIN THAY ĐỔI SAU NÀY")
    _format_run(
        run,
        style.body_font,
        style.notes_font_size_pt,
        bold=True,
    )
    note = cell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(note, line_pt=14)
    run = note.add_run("Bản nháp để rà soát - không ghi thay đổi hộ tịch tại đây")
    _format_run(run, style.body_font, 9, italic=True, color="777777")


def _new_fixed_table(
    doc: DocxDocument,
    widths_mm: list[float],
    *,
    hidden_borders: bool,
    total_width_mm: float | None = None,
) -> Table:
    table = doc.add_table(rows=1, cols=len(widths_mm))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_geometry(table, widths_mm, total_width_mm=total_width_mm)
    _set_table_borders(table, hidden=hidden_borders)
    for row in table.rows:
        _prevent_row_split(row._tr)
        for cell in row.cells:
            _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def _set_table_geometry(
    table: Table,
    widths_mm: list[float],
    *,
    total_width_mm: float | None,
) -> None:
    widths_dxa = [int(Mm(width).twips) for width in widths_mm]
    total_dxa = int(Mm(total_width_mm).twips) if total_width_mm else sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for column, width_mm, width_dxa in zip(
        table.columns, widths_mm, widths_dxa, strict=True
    ):
        column.width = Mm(width_mm)
        for cell in column.cells:
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_dxa))


def _set_table_borders(table: Table, *, hidden: bool) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil" if hidden else "single")
        if not hidden:
            element.set(qn("w:sz"), "8")
            element.set(qn("w:color"), "777777")


def _set_cell_margins(
    cell: _Cell, *, top: int, start: int, bottom: int, end: int
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = tc_mar.find(qn(f"w:{name}"))
        if margin is None:
            margin = OxmlElement(f"w:{name}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _prevent_row_split(row_element) -> None:
    tr_pr = row_element.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_cell_line(
    cell: _Cell,
    label: str,
    value: str,
    template: DraftTemplate,
    *,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    _format_body_paragraph(paragraph, template)
    _append_label_value(paragraph, label, value, template)


def _append_label_value(
    paragraph: Paragraph,
    label: str,
    value: str,
    template: DraftTemplate,
    *,
    value_bold: bool = False,
) -> None:
    label_run = paragraph.add_run(f"{label}: ")
    _format_run(
        label_run,
        template.docx_style.body_font,
        template.docx_style.body_size_pt,
        bold=True,
    )
    value_run = paragraph.add_run(value)
    missing = value.startswith("[CHƯA CÓ:")
    _format_run(
        value_run,
        template.docx_style.body_font,
        template.docx_style.body_size_pt,
        bold=value_bold,
        italic=missing,
        color="777777" if missing else None,
    )


def _format_body_paragraph(paragraph: Paragraph, template: DraftTemplate) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(template.docx_style.line_spacing_pt)


def _set_paragraph_spacing(paragraph: Paragraph, *, line_pt: float) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(line_pt)


def _format_run(
    run: Run,
    font_name: str,
    size_pt: float,
    *,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _set_style_fonts(style_element, font_name: str) -> None:
    r_pr = style_element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _value(values: dict[str, str], key: str, label: str) -> str:
    return values.get(key) or f"[CHƯA CÓ: {label}]"


def _date_value(
    values: dict[str, str], key: str, label: str, *, words: bool
) -> str:
    raw = values.get(key)
    if not raw:
        return f"[CHƯA CÓ: {label}]"
    parsed = date.fromisoformat(raw)
    if not words:
        return parsed.strftime("%d/%m/%Y")
    from src.services.drafts.renderer import _date_in_words

    return _date_in_words(parsed)


_BUILDERS: dict[
    str, Callable[[DraftTemplate, GeneratedDraft], DocxDocument]
] = {
    "khai_sinh.giay_khai_sinh": _build_birth_certificate,
}
