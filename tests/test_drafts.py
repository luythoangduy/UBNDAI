import importlib.util
import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from fastapi import FastAPI
from fastapi.testclient import TestClient
import pytest

from src.models import DraftGenerateRequest, DraftTemplateImportRequest
from src.services import drafts
from src.services.drafts import registry as draft_registry
from src.services.drafts import template_importer
from src.services.ocr.engine import OcrField, OcrResult, VisionLlmEngine
from src.services.procedure_capabilities import capabilities_for, form_schema_for


def _draft_api_client() -> TestClient:
    """Load riêng router để test không phụ thuộc các workstream đang merge."""

    path = Path(__file__).parents[1] / "src" / "api" / "v1" / "drafts.py"
    spec = importlib.util.spec_from_file_location("draft_api_under_test", path)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    app = FastAPI()
    app.include_router(module.router, prefix="/api/v1")
    return TestClient(app)


def _birth_certificate_values() -> dict:
    return {
        "ho_ten_con": "Nguyễn An",
        "ngay_sinh": "2026-07-17",
        "gioi_tinh": "Nữ",
        "dan_toc": "Kinh",
        "quoc_tich": "Việt Nam",
        "noi_sinh": "Bệnh viện A, phường X, thành phố Hà Nội",
        "que_quan": "phường X, thành phố Hà Nội",
        "so_dinh_danh_ca_nhan": "012345678901",
        "ho_ten_me": "Nguyễn Thị B",
        "nam_sinh_me": "1995",
        "dan_toc_me": "Kinh",
        "quoc_tich_me": "Việt Nam",
        "noi_cu_tru_me": "phường X, thành phố Hà Nội",
        "noi_dang_ky": "Ủy ban nhân dân phường X, thành phố Hà Nội",
        "ngay_dang_ky": "2026-07-18",
    }


@pytest.fixture(autouse=True)
def clear_draft_cache():
    drafts.clear_cache()
    yield
    drafts.clear_cache()


def test_birth_certificate_uses_procedure_specific_legal_template():
    template = drafts.get_template("khai_sinh")

    assert template.id == "khai_sinh.giay_khai_sinh"
    assert any(
        source.document_number == "04/2020/TT-BTP"
        and source.role == "output_template"
        for source in template.legal_sources
    )
    assert all(
        source.document_number != "30/2020/NĐ-CP"
        for source in template.legal_sources
    )


def test_generate_birth_certificate_normalizes_rules_from_circular():
    result = drafts.generate(
        DraftGenerateRequest(
            procedure_id="khai_sinh", values=_birth_certificate_values()
        )
    )

    assert result.status == "draft"
    assert result.ready_for_review is True
    assert result.normalized_values["ho_ten_con"] == "NGUYỄN AN"
    assert "17/07/2026" in result.rendered_text
    assert "Ngày mười bảy tháng bảy năm hai nghìn không trăm hai mươi sáu" in result.rendered_text
    assert result.rendered_text.startswith("DỰ THẢO - KHÔNG CÓ GIÁ TRỊ PHÁP LÝ")
    assert "chữ ký" in result.disclaimer.casefold()


def test_missing_required_data_is_rejected_by_default():
    values = _birth_certificate_values()
    del values["noi_sinh"]

    with pytest.raises(drafts.DraftDataError) as exc_info:
        drafts.generate(
            DraftGenerateRequest(procedure_id="khai_sinh", values=values)
        )

    assert exc_info.value.fields == ["noi_sinh"]


def test_partial_father_information_requires_complete_group():
    values = _birth_certificate_values() | {"ho_ten_cha": "Trần Văn C"}

    with pytest.raises(drafts.DraftDataError) as exc_info:
        drafts.generate(
            DraftGenerateRequest(procedure_id="khai_sinh", values=values)
        )

    assert set(exc_info.value.fields) == {
        "nam_sinh_cha",
        "dan_toc_cha",
        "quoc_tich_cha",
        "noi_cu_tru_cha",
    }


def test_allow_incomplete_returns_explicit_placeholders():
    result = drafts.generate(
        DraftGenerateRequest(
            procedure_id="khai_sinh",
            values={"ho_ten_con": "Nguyễn An"},
            allow_incomplete=True,
        )
    )

    assert result.ready_for_review is False
    assert "ngay_sinh" in result.missing_required_fields
    assert "[CHƯA CÓ: Ngày, tháng, năm sinh (bằng số)]" in result.rendered_text


def test_draft_api_exposes_template_and_generates_preview():
    client = _draft_api_client()

    templates = client.get("/api/v1/drafts/templates/khai_sinh")
    assert templates.status_code == 200
    assert templates.json()[0]["legal_sources"][0]["document_number"] == "04/2020/TT-BTP"

    response = client.post(
        "/api/v1/drafts/generate",
        json={"procedure_id": "khai_sinh", "values": _birth_certificate_values()},
    )
    assert response.status_code == 200
    assert response.json()["output_name"] == "Giấy khai sinh"


def test_draft_api_returns_clear_error_when_template_is_missing():
    client = _draft_api_client()

    response = client.post(
        "/api/v1/drafts/generate",
        json={"procedure_id": "unknown", "values": {}},
    )

    assert response.status_code == 404


def test_birth_certificate_docx_encodes_regulated_page_and_typography():
    generated = drafts.generate_docx(
        DraftGenerateRequest(
            procedure_id="khai_sinh", values=_birth_certificate_values()
        )
    )

    assert generated.filename == "giay-khai-sinh-du-thao.docx"
    assert generated.content.startswith(b"PK")
    document = Document(BytesIO(generated.content))
    section = document.sections[0]
    assert section.page_width.mm == pytest.approx(210, abs=0.1)
    assert section.page_height.mm == pytest.approx(297, abs=0.1)
    assert section.top_margin.mm == pytest.approx(12.8, abs=0.1)
    assert section.right_margin.mm == pytest.approx(23, abs=0.1)
    assert section.bottom_margin.mm == pytest.approx(18, abs=0.1)
    assert section.left_margin.mm == pytest.approx(23, abs=0.1)

    normal = document.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size.pt == pytest.approx(13)
    assert normal.paragraph_format.line_spacing.pt == pytest.approx(21.5)
    style_fonts = normal._element.rPr.rFonts
    assert style_fonts.get(qn("w:ascii")) == "Times New Roman"
    assert style_fonts.get(qn("w:hAnsi")) == "Times New Roman"
    assert style_fonts.get(qn("w:eastAsia")) == "Times New Roman"

    title = next(p for p in document.paragraphs if p.text == "GIẤY KHAI SINH")
    assert title.runs[0].font.size.pt == pytest.approx(22)
    assert title.runs[0].bold is True
    assert str(title.runs[0].font.color.rgb) == "C00000"
    assert "DỰ THẢO - KHÔNG CÓ GIÁ TRỊ PHÁP LÝ" in section.header.paragraphs[0].text
    document_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "Nguyễn Thị B" in document_text

    notes_cell = document.tables[-1].cell(0, 0)
    assert "PHẦN GHI CHÚ NHỮNG THÔNG TIN THAY ĐỔI SAU NÀY" in notes_cell.text
    assert document.tables[-1].rows[0].height.mm == pytest.approx(260, abs=0.1)
    assert len(document.sections) == 2
    assert document.sections[1].left_margin.mm == pytest.approx(26, abs=0.1)
    assert document.sections[1].right_margin.mm == pytest.approx(26, abs=0.1)


def test_draft_docx_api_returns_downloadable_word_file():
    client = _draft_api_client()
    response = client.post(
        "/api/v1/drafts/generate.docx",
        json={"procedure_id": "khai_sinh", "values": _birth_certificate_values()},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == drafts.DOCX_MEDIA_TYPE
    assert response.headers["content-disposition"] == (
        'attachment; filename="giay-khai-sinh-du-thao.docx"'
    )
    assert response.headers["x-draft-legal-status"] == "review-only"
    assert response.content.startswith(b"PK")


def test_template_store_has_to_khai_for_citizen_request():
    """Kho template phải có tờ khai (văn bản người dân gửi cán bộ) bên cạnh giấy khai sinh."""
    templates = drafts.list_templates("khai_sinh")
    ids = {template.id for template in templates}
    assert "khai_sinh.to_khai" in ids
    to_khai = drafts.get_template("khai_sinh", "khai_sinh.to_khai")
    assert to_khai.is_default is False
    assert to_khai.docx_style.filename == "to-khai-dang-ky-khai-sinh.docx"
    assert any(s.role == "output_template" for s in to_khai.legal_sources)
    # Mặc định của procedure vẫn là giấy khai sinh (kết quả cho cán bộ).
    assert drafts.get_template("khai_sinh").id == "khai_sinh.giay_khai_sinh"


def test_export_docx_from_editor_html():
    """POST /drafts/export.docx: DOCX chứa đúng nội dung HTML editor, font Times New Roman."""
    client = _draft_api_client()
    html = (
        "<h2>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM<br/>Độc lập - Tự do - Hạnh phúc</h2>"
        "<h1>TỜ KHAI ĐĂNG KÝ KHAI SINH</h1>"
        "<p>Họ, chữ đệm, tên: <strong>NGUYỄN VĂN AN</strong></p>"
        "<p><em>Tôi cam đoan nội dung khai trên đây là đúng sự thật.</em></p>"
    )
    response = client.post(
        "/api/v1/drafts/export.docx",
        json={"html": html, "filename": "to-khai-dang-ky-khai-sinh.docx"},
    )
    assert response.status_code == 200, response.text
    assert "to-khai-dang-ky-khai-sinh.docx" in response.headers["content-disposition"]

    document = Document(BytesIO(response.content))
    texts = [p.text for p in document.paragraphs if p.text.strip()]
    assert any("TỜ KHAI ĐĂNG KÝ KHAI SINH" in t for t in texts)
    assert any("NGUYỄN VĂN AN" in t for t in texts)
    bold_runs = [
        run.text
        for p in document.paragraphs
        for run in p.runs
        if run.bold and run.text.strip()
    ]
    assert "NGUYỄN VĂN AN" in bold_runs
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            r_fonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
            assert r_fonts is not None
            assert r_fonts.get(qn("w:ascii")) == "Times New Roman"


def test_export_docx_rejects_bad_filename():
    client = _draft_api_client()
    response = client.post(
        "/api/v1/drafts/export.docx",
        json={"html": "<p>x</p>", "filename": "../../evil.docx"},
    )
    assert response.status_code == 422


def test_non_catalog_template_can_generate_preview_and_generic_docx(
    tmp_path, monkeypatch
):
    """A source-backed manifest must work without a catalog procedure record."""
    source = Path(__file__).parents[1] / "data" / "draft_templates" / "khai_sinh_to_khai.json"
    manifest = json.loads(source.read_text(encoding="utf-8"))
    manifest.update(
        {
            "id": "searched-procedure.request-form",
            "procedure_id": "searched-procedure",
            "output_name": "Mẫu tìm từ nguồn chính thức",
            "is_default": True,
        }
    )
    (tmp_path / "searched.json").write_text(
        json.dumps(manifest, ensure_ascii=False), encoding="utf-8"
    )
    monkeypatch.setattr(draft_registry, "DEFAULT_TEMPLATE_DIR", tmp_path)
    drafts.clear_cache()

    result = drafts.generate(
        DraftGenerateRequest(
            procedure_id="searched-procedure",
            values={},
            allow_incomplete=True,
        )
    )
    document = drafts.generate_docx(
        DraftGenerateRequest(
            procedure_id="searched-procedure",
            values={},
            allow_incomplete=True,
        )
    )
    caps = capabilities_for("searched-procedure")
    schema = form_schema_for("searched-procedure")

    assert result.template_id == "searched-procedure.request-form"
    assert result.ready_for_review is False
    assert result.legal_sources
    assert document.content.startswith(b"PK")
    assert caps.dynamic_form is True
    assert caps.ocr_autofill is True
    assert caps.requires_human_review is True
    assert schema is not None and schema.fields


async def test_official_searched_template_is_ocr_imported_without_catalog(
    monkeypatch,
):
    async def fake_download(url: str) -> tuple[bytes, str]:
        return b"\xff\xd8\xff searched-template", "image/jpeg"

    engine = VisionLlmEngine(provider="openai", api_key="test", model="test")
    monkeypatch.setattr(template_importer, "_download", fake_download)
    monkeypatch.setattr(template_importer, "get_engine", lambda: engine)
    monkeypatch.setattr(
        engine,
        "extract",
        lambda image, field_keys, task: OcrResult(
            raw_text="Họ tên: ........ Ngày sinh: ../../....",
            fields=[
                OcrField(
                    key="ho_ten",
                    value="__TEXT__",
                    confidence=0.96,
                    note="Họ và tên",
                ),
                OcrField(
                    key="ngay_sinh",
                    value="__DATE__",
                    confidence=0.94,
                    note="Ngày sinh",
                ),
            ],
            engine="vision_llm",
        ),
    )

    template = await template_importer.import_template(
        DraftTemplateImportRequest(
            procedure_id="searched-live",
            source_url="https://example.gov.vn/mau.pdf",
            title="Mẫu trực tuyến",
        )
    )

    assert template.procedure_id == "searched-live"
    assert [field.key for field in template.fields] == ["ho_ten", "ngay_sinh"]
    assert template.fields[1].input_type == "date"
    assert template.fields[0].required is False
    assert "đối chiếu" in template.disclaimer.casefold()
    assert drafts.get_template("searched-live", template.id).id == template.id


def test_searched_template_import_rejects_non_official_source():
    with pytest.raises(template_importer.TemplateImportError):
        template_importer._validate_source_url("https://example.com/form.pdf")
