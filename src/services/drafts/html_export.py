"""Xuất DOCX từ HTML editor — port cách làm của C2-App-108 (services/drafting.py).

Người dân soạn/sửa tờ khai trong editor (Tiptap) rồi tải xuống: file DOCX phải
chứa đúng nội dung đang thấy, nên parse trực tiếp HTML của editor thay vì render
lại từ values qua builder cứng. Hỗ trợ block p / h1-h6 / ul / ol và inline
b / i / u / strong / em / br / span (font-size), căn giữa h1-h2 theo thể thức
văn bản hành chính (Nghị định 30/2020/NĐ-CP dùng Times New Roman).
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from src.models import DraftHtmlExportRequest

DOCX_FONT_NAME = "Times New Roman"

_BLOCK_PATTERN = re.compile(r"<(p|h[1-6]|ul|ol)[^>]*>(.*?)</\1>", re.IGNORECASE)
_LI_PATTERN = re.compile(r"<li[^>]*>(.*?)</li>", re.IGNORECASE)
_INLINE_SPLIT = re.compile(r"(</?(?:b|i|u|strong|em|br|span)[^>]*>)", re.IGNORECASE)


def _apply_font(run) -> None:
    """Ép Times New Roman cả slot eastAsia để dấu tiếng Việt cùng mặt chữ."""
    run.font.name = DOCX_FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), DOCX_FONT_NAME)


def _extract_font_size_pt(tag: str) -> float | None:
    match = re.search(r"font-size\s*:\s*([0-9]+(?:\.[0-9]+)?)(pt|px)?", tag, re.IGNORECASE)
    if not match:
        return None
    value = float(match.group(1))
    unit = (match.group(2) or "pt").lower()
    return value * 0.75 if unit == "px" else value


def _parse_runs(paragraph, html_content: str) -> None:
    clean_html = re.sub(r"</?(div|a|img)[^>]*>", "", html_content)
    tokens = _INLINE_SPLIT.split(clean_html)

    bold = False
    italic = False
    underline = False
    font_size_stack: list[float | None] = []

    for token in tokens:
        token_lower = token.lower()
        if token_lower in ("<b>", "<strong>"):
            bold = True
        elif token_lower in ("</b>", "</strong>"):
            bold = False
        elif token_lower in ("<i>", "<em>"):
            italic = True
        elif token_lower in ("</i>", "</em>"):
            italic = False
        elif token_lower == "<u>":
            underline = True
        elif token_lower == "</u>":
            underline = False
        elif token_lower.startswith("<span"):
            font_size_stack.append(_extract_font_size_pt(token))
        elif token_lower == "</span>":
            if font_size_stack:
                font_size_stack.pop()
        elif token_lower in ("<br>", "<br/>", "<br />"):
            paragraph.add_run("\n")
        elif not token.startswith("<"):
            if token:
                text = (
                    token.replace("&nbsp;", " ")
                    .replace("&lt;", "<")
                    .replace("&gt;", ">")
                    .replace("&quot;", '"')
                    .replace("&amp;", "&")
                )
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                active_font_size = next(
                    (size for size in reversed(font_size_stack) if size), None
                )
                if active_font_size:
                    run.font.size = Pt(active_font_size)


def export_html_docx(request: DraftHtmlExportRequest) -> bytes:
    doc = Document()

    normal_font = doc.styles["Normal"].font
    normal_font.name = DOCX_FONT_NAME
    normal_font.size = Pt(13)
    normal_rpr = doc.styles["Normal"].element.get_or_add_rPr()
    normal_rfonts = normal_rpr.find(qn("w:rFonts"))
    if normal_rfonts is None:
        normal_rfonts = normal_rpr.makeelement(qn("w:rFonts"), {})
        normal_rpr.append(normal_rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        normal_rfonts.set(qn(attr), DOCX_FONT_NAME)

    if request.title:
        doc.add_heading(request.title, 0)

    html = request.html.replace("\n", " ")
    for match in _BLOCK_PATTERN.finditer(html):
        tag = match.group(1).lower()
        content = match.group(2)

        if tag.startswith("h"):
            level = int(tag[1])
            paragraph = doc.add_heading("", level=level)
            # Quốc hiệu / tên loại văn bản đặt ở h1-h2 — căn giữa theo thể thức.
            if level <= 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == "p":
            paragraph = doc.add_paragraph()
            opening = match.group(0)
            if 'text-align="center"' in opening or "text-align: center" in opening:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align="right"' in opening or "text-align: right" in opening:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align="justify"' in opening or "text-align: justify" in opening:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:  # ul | ol
            style = "List Bullet" if tag == "ul" else "List Number"
            for li_match in _LI_PATTERN.finditer(content):
                paragraph = doc.add_paragraph(style=style)
                _parse_runs(paragraph, li_match.group(1))
            continue

        _parse_runs(paragraph, content)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _apply_font(run)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
